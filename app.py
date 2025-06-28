# app.py
from flask import Flask, render_template, request, jsonify, send_file, send_from_directory
from docx import Document
import pandas as pd
from io import BytesIO
import os
import re
import time
from datetime import datetime
import json
import uuid
from functools import wraps
from flask_compress import Compress  # 新增压缩支持

# 初始化Flask应用
app = Flask(__name__)
Compress(app)  # 启用Gzip压缩
app.config['UPLOAD_FOLDER'] = 'temp_uploads'
app.config['OUTPUT_FOLDER'] = 'processed_files'
app.config['HISTORY_FILE'] = 'history.json'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB单文件限制
app.config['BATCH_MAX_SIZE'] = 100 * 1024 * 1024    # 班评模式100MB总限制
app.config['ADMIN_PASSWORD'] = os.getenv('ADMIN_PASS', 'ADMIN123')  # 管理员密码
# 新增自评统计信息存储
app.config['SELF_TOTAL_FILES'] = 0
app.config['SELF_TOTAL_TIME'] = 0
# 新增班评统计信息存储
app.config['BATCH_TOTAL_FILES'] = 0
app.config['BATCH_TOTAL_TIME'] = 0

# 确保目录存在
for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
    if not os.path.exists(folder):
        os.makedirs(folder)

# 智能频率限制装饰器
def smart_rate_limit():
    def decorator(f):
        self_times = []
        @wraps(f)
        def wrapper(*args, **kwargs):
            evaluation_type = request.form.get('evaluationType', 'self')

            # 自评模式限速
            if evaluation_type == 'self':
                now = time.time()
                self_times.append(now)
                self_times[:] = [t for t in self_times if t > now - 60]
                if len(self_times) > 3:  # 每分钟最多3次
                    return jsonify({"success": False, "error": "操作过于频繁，请1分钟后再试"}), 429
            return f(*args, **kwargs)
        return wrapper
    return decorator

def load_history():
    if os.path.exists(app.config['HISTORY_FILE']):
        try:
            with open(app.config['HISTORY_FILE'], 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"读取历史记录出错: {str(e)}")
    return []

def save_history(history):
    try:
        # 超过 5 条记录时自动删除旧记录
        while len(history) > 5:
            oldest_entry = history.pop()
            if os.path.exists(oldest_entry['file_path']):
                os.remove(oldest_entry['file_path'])

        with open(app.config['HISTORY_FILE'], 'w', encoding='utf-8') as f:
            json.dump(history, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"保存历史记录出错: {str(e)}")

def parse_word_doc(file_stream, evaluation_type):
    doc = Document(file_stream)
    student_data = {
        "学号": "未提取",
        "姓名": "未提取",
        "思想品德": 0.0,
        "专业科研": 0.0,
        "体艺": 0.0,
        "劳动实践": 0.0
    }

    try:
        # 提取学号姓名
        for para in doc.paragraphs:
            text = para.text.strip()
            id_match = re.search(r'学号[:：]\s*(\d+)', text)
            name_match = re.search(r'姓名[:：]\s*(\S+)', text)
            if id_match:
                student_data["学号"] = id_match.group(1)
            if name_match:
                student_data["姓名"] = name_match.group(1)

        # 分数统计逻辑
        current_category = None
        category_mapping = {
            "品德": "思想品德",
            "专业与科研": "专业科研",
            "体艺": "体艺",
            "劳动与实践": "劳动实践"
        }

        for table in doc.tables:
            for row in table.rows:
                if any("项目" in cell.text for cell in row.cells):
                    continue

                for cell in row.cells:
                    text = cell.text.strip().split('(')[0].strip()
                    if text in category_mapping:
                        current_category = category_mapping[text]
                        break

                # 根据评估类型选择列索引
                if evaluation_type == 'batch':
                    column_index = 4  # 班评模式第5列
                else:
                    column_index = 5  # 自评模式第6列

                if current_category and len(row.cells) > column_index:
                    cell_text = row.cells[column_index].text.strip()
                    if cell_text:
                        try:
                            value = float(cell_text)
                            student_data[current_category] += value
                        except ValueError:
                            pass

        return student_data

    except Exception as e:
        print(f"解析异常：{str(e)}")
        raise ValueError(f"文档解析失败: {str(e)}")

@app.route('/')
def index():
    statistics = calculate_statistics()
    return render_template('index.html', statistics=statistics)

@app.route('/help')
def help_page():
    return render_template('help.html')

@app.route('/share')  # 这里定义 share_page 对应的路由
def share_page():
    return render_template('share.html')

@app.route('/study')
def study_page():
    return render_template('study.html')

@app.route('/process', methods=['POST'])
@smart_rate_limit()
def process_files():
    try:
        start_time = time.time()
        evaluation_type = request.form.get('evaluationType', 'self')
        batch_files = request.files.getlist('batchFiles')
        self_files = request.files.getlist('selfFiles')

        # 班评模式总大小校验
        if evaluation_type == 'batch':
            total_size = sum(len(file.read()) for file in batch_files)
            if total_size > app.config['BATCH_MAX_SIZE']:
                return jsonify({
                    "success": False,
                    "error": f"总文件大小超过限制（{app.config['BATCH_MAX_SIZE']//1024//1024}MB）"
                }), 400
            for file in batch_files:
                file.seek(0)

        processed_data = []
        file_names = []

        for file in batch_files + self_files:
            if file.filename == '':
                continue

            file_names.append(file.filename)
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(temp_path)

            try:
                with open(temp_path, 'rb') as f:
                    data = parse_word_doc(f, evaluation_type)
                    processed_data.append(data)
            except Exception as e:
                print(f"解析文件 {file.filename} 时出错: {str(e)}")
                os.remove(temp_path)
                return jsonify({"success": False, "error": f"解析文件 {file.filename} 时出错: {str(e)}"})

            os.remove(temp_path)

        if not processed_data:
            return jsonify({"success": False, "error": "没有可处理的文件"})

        df = pd.DataFrame(processed_data)
        excel_buffer = BytesIO()

        try:
            with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='总评分', index=False)
                df[['学号', '姓名', '思想品德']].to_excel(writer, sheet_name='思想品德', index=False)
                df[['学号', '姓名', '专业科研']].to_excel(writer, sheet_name='专业科研', index=False)
                df[['学号', '姓名', '体艺']].to_excel(writer, sheet_name='体艺', index=False)
                df[['学号', '姓名', '劳动实践']].to_excel(writer, sheet_name='劳动实践', index=False)
        except Exception as e:
            print(f"生成Excel文件时出错: {str(e)}")
            return jsonify({"success": False, "error": f"生成Excel文件时出错: {str(e)}"})

        file_uuid = str(uuid.uuid4())
        output_filename = f"综合测评结果_{datetime.now().strftime('%Y%m%d%H%M%S')}_{file_uuid}.xlsx"
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)

        with open(output_path, 'wb') as f:
            f.write(excel_buffer.getvalue())

        processing_time = time.time() - start_time

        if evaluation_type == 'batch':
            history = load_history()
            history_entry = {
                'id': file_uuid,
                'type': evaluation_type,
                'file_name': '综合测评结果.xlsx',
                'original_files': file_names,
                'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'processing_time': processing_time,
                'file_path': output_path
            }
            history.insert(0, history_entry)
            save_history(history)
            app.config['BATCH_TOTAL_FILES'] += len(batch_files)
            app.config['BATCH_TOTAL_TIME'] += processing_time
        else:
            # 更新自评统计信息
            app.config['SELF_TOTAL_FILES'] += len(self_files)
            app.config['SELF_TOTAL_TIME'] += processing_time

        return jsonify({
            "success": True,
            "html_table": df.to_html(
                classes="table table-striped table-hover",
                index=False,
                border=0),
            "excel_url": f"/download/{file_uuid}"
        })
    except Exception as e:
        print(f"处理请求时出错: {str(e)}")
        return jsonify({"success": False, "error": str(e)})

@app.route('/download/<file_id>')
def download_excel(file_id):
    password = request.args.get('password')
    history = load_history()
    for entry in history:
        if entry['id'] == file_id:
            if entry['type'] == 'self':
                # 自评模式无需密码
                if os.path.exists(entry['file_path']):
                    return send_file(
                        entry['file_path'],
                        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                        as_attachment=True,
                        download_name=entry['file_name']
                    )
                else:
                    return jsonify({"success": False, "error": "文件不存在"}), 404
            elif entry['type'] == 'batch':
                # 班评模式需要密码
                if password != app.config['ADMIN_PASSWORD']:
                    return jsonify({"success": False, "error": "需要管理员密码"}), 401
                if os.path.exists(entry['file_path']):
                    return send_file(
                        entry['file_path'],
                        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                        as_attachment=True,
                        download_name=entry['file_name']
                    )
                else:
                    return jsonify({"success": False, "error": "文件不存在"}), 404

    return jsonify({"success": False, "error": "找不到该文件"}), 404

@app.route('/history')
def get_history():
    try:
        history = load_history()
        return jsonify(history[:5])
    except Exception as e:
        print(f"获取历史记录时出错: {str(e)}")
        return jsonify([])

@app.route('/favicon.ico')
def favicon():
    return send_from_directory(
        os.path.join(app.root_path, 'static'),
        'favicon.ico',
        mimetype='image/vnd.microsoft.icon'
    )

def calculate_statistics():
    try:
        total_files = app.config['SELF_TOTAL_FILES'] + app.config['BATCH_TOTAL_FILES']

        if total_files == 0:
            return {
                "全站处理文件总数": 0,
                "全站平均处理时间(秒)": 0
            }

        total_time = app.config['SELF_TOTAL_TIME'] + app.config['BATCH_TOTAL_TIME']
        avg_time = total_time / total_files

        return {
            "全站处理文件总数": total_files,
            "全站平均处理时间(秒)": round(avg_time, 2)
        }
    except Exception as e:
        print(f"计算统计信息时出错: {str(e)}")
        return {
            "全站处理文件总数": "计算失败",
            "全站平均处理时间(秒)": "计算失败"
        }

@app.route('/delete/<file_id>', methods=['DELETE'])
def delete_history_file(file_id):
    password = request.args.get('password')
    if password != app.config['ADMIN_PASSWORD']:
        return jsonify({"success": False, "error": "需要管理员密码"}), 401

    history = load_history()
    new_history = []
    for entry in history:
        if entry['id'] == file_id:
            if os.path.exists(entry['file_path']):
                os.remove(entry['file_path'])
        else:
            new_history.append(entry)

    save_history(new_history)
    return jsonify({"success": True, "message": "删除成功"})

# 添加缓存控制头
@app.after_request
def add_cache_header(response):
    if request.path.startswith('/static'):
        response.headers['Cache-Control'] = 'public, max-age=86400'
    return response

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=80, debug=True)
